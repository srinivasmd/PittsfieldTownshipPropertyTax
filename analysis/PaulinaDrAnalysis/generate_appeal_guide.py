#!/usr/bin/env python3
"""Generate Property Tax Appeal Guide as Word template for neighbors."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

OUTPUT = os.path.join(os.path.dirname(__file__), "Property_Tax_Appeal_Guide_Paulina_Dr.docx")

def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)
    return h

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = val
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.size = Pt(9)
    return table

def main():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title
    title = doc.add_heading('2026 Property Tax Appeal Guide', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Meadows of Arbor Ridge (AR-4) | Pittsfield Township')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(100, 100, 100)

    # Deadline box
    deadline = doc.add_paragraph()
    deadline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = deadline.add_run('\nDEADLINE: Written appeals must be received by March 10, 2026 at 5:00 PM\n')
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(180, 0, 0)
    run2 = deadline.add_run('Pittsfield Township Assessing, 6201 W. Michigan Ave, Ann Arbor, MI 48108\n')
    run2.font.size = Pt(10)

    doc.add_paragraph('_' * 75)

    # Section 1: Why Appeal
    add_heading_styled(doc, 'Why You Should Consider Appealing', level=1)

    doc.add_paragraph(
        'The 2026 assessments for homes in the Meadows of Arbor Ridge subdivision (AR-4) '
        'appear to significantly overstate market value. The township\'s own data confirms this:'
    )

    bullets = [
        'The Economic Condition Factor (ECF) for AR-4 is 0.802, meaning the township\'s '
        'cost-approach values exceed actual sale prices by approximately 20%.',
        'No home in AR-4 has ever resold at or above $485,000. The highest resale was $500,000 '
        '(4711 Paulina Dr, a builder first-sale in May 2023).',
        'The average arm\'s-length sale price across 10+ sales in AR-4 is approximately $459,690.',
        'ECF has been consistently below 1.0 for three years: 0.744 (2024) -> 0.790 (2025) -> 0.802 (2026).',
        'Land values have increased 17.7% in 3 years: $84,300 (2023) -> $89,800 (2024) -> $96,400 (2025) -> $99,200 (2026).',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    # Section 2: How Assessment Works
    add_heading_styled(doc, 'How Michigan Property Tax Assessment Works', level=1)
    doc.add_paragraph(
        'Under Michigan law, the Assessed Value (SEV) must equal 50% of the property\'s True Cash '
        'Value (TCV), defined as the "usual selling price" - the price a willing buyer and seller '
        'would agree to in an arm\'s-length transaction (MCL 211.27). Your Taxable Value (TV) is '
        'capped by Proposal A at the lesser of 5% or the CPI increase per year.'
    )
    doc.add_paragraph(
        'Even if reducing the assessed value does not immediately lower your tax bill (because '
        'TV < AV), it constrains future taxable value growth and establishes a lower baseline '
        'for when the property transfers or when TV catches up to AV.'
    )

    # Section 3: Template section for individual use
    add_heading_styled(doc, 'Your Property Information (Fill In)', level=1)

    p = doc.add_paragraph()
    p.add_run('Parcel Number: ').bold = True
    p.add_run('L-12-13-___-___    ')
    p.add_run('(find on your tax bill or BSA Online)')
    p.runs[-1].font.color.rgb = RGBColor(128, 128, 128)

    p2 = doc.add_paragraph()
    p2.add_run('Property Address: ').bold = True
    p2.add_run('________________________________________')

    p3 = doc.add_paragraph()
    p3.add_run('2026 Assessed Value (SEV): ').bold = True
    p3.add_run('$___________    ')
    p3.add_run('Implied TCV (SEV x 2): ').bold = True
    p3.add_run('$___________')

    p4 = doc.add_paragraph()
    p4.add_run('2026 Taxable Value: ').bold = True
    p4.add_run('$___________')

    p5 = doc.add_paragraph()
    p5.add_run('Your Requested TCV: ').bold = True
    p5.add_run('$___________    ')
    p5.add_run('Your Requested SEV: ').bold = True
    p5.add_run('$___________')

    doc.add_paragraph()

    # Section 4: Comparable Sales
    add_heading_styled(doc, 'Comparable Sales in Our Subdivision (AR-4)', level=1)

    doc.add_paragraph(
        'The following arm\'s-length sales have occurred in Meadows of Arbor Ridge. '
        'All ECF values below 1.0 confirm the cost approach OVERVALUES these properties.'
    )

    sales_headers = ["Address", "Sale Price", "Date", "ECF '26", "ECF '25", "ECF '24"]
    sales_data = [
        ["4807 Paulina Dr", "$450,000", "07/2023", "0.766", "0.764", "--"],
        ["4918 Paulina Dr", "$447,000", "03/2023", "--", "0.803", "0.804"],
        ["4711 Paulina Dr", "$500,000", "05/2023", "0.800", "0.797", "--"],
        ["4782 Paulina Dr", "$452,000", "10/2022", "--", "0.806", "0.807"],
        ["4930 Paulina Dr", "$365,424", "07/2021", "--", "--", "0.707"],
        ["4610 Lilac Lane", "$465,000", "08/2023", "0.822", "0.819", "--"],
        ["4617 Lilac Lane", "$452,000", "06/2024", "--", "--", "--"],
        ["4324 Christina Ct", "$499,900", "08/2023", "0.818", "0.815", "0.643"],
        ["4321 Christina Ct", "$430,000", "05/2022", "--", "0.723", "0.724"],
        ["4470 Connor Dr", "$481,000", "05/2022", "--", "0.792", "0.799"],
        ["4458 Christina Dr", "$420,000", "01/2022", "--", "--", "0.724"],
    ]
    add_table(doc, sales_headers, sales_data)
    doc.add_paragraph()

    # Adjacent comps
    add_heading_styled(doc, 'Adjacent Subdivision Sales (AR-1, AR-2)', level=2)
    adj_headers = ["Address (Area)", "Sale Price", "Date", "Cur. Appraisal"]
    adj_data = [
        ["4562 Christina Dr (AR-1)", "$427,000", "11/2024", "$389,440"],
        ["4936 Matthew Ct (AR-1)", "$390,000", "10/2024", "$395,088"],
        ["4579 Connor Ct (AR-1)", "$386,000", "02/2024", "$413,397"],
        ["4336 Cloverlane Dr (AR-2)", "$455,000", "11/2024", "$398,249"],
        ["4236 Cloverlane Dr (AR-2)", "$410,000", "05/2023", "$437,059"],
        ["4249 Cloverlane Dr (AR-2)", "$430,000", "06/2023", "$424,507"],
    ]
    add_table(doc, adj_headers, adj_data)
    doc.add_paragraph()

    # Section 5: Filing Instructions
    add_heading_styled(doc, 'How to File Your Appeal', level=1)

    add_heading_styled(doc, 'Step 1: Board of Review (Required First Step)', level=2)
    doc.add_paragraph(
        'You MUST file with the Pittsfield Township Board of Review first. '
        'Skipping this forfeits your right to appeal to the Michigan Tax Tribunal.'
    )
    step1_bullets = [
        'Public Hearings: March 9 (9am-12pm and 6pm-9pm), March 10 (9am-12pm), March 11 (1pm-5pm)',
        'Written Appeal Deadline: March 10, 2026 at 5:00 PM',
        'Location: 6201 W. Michigan Avenue, Ann Arbor, MI 48108',
        'Phone: 734-822-3115 | Email: assessing@pittsfield-mi.gov',
        'Form Required: Petition to Board of Review (Form L-4035)',
        'Written petitions are accepted in lieu of personal appearance',
        'If mailing, send certified with return receipt requested',
    ]
    for b in step1_bullets:
        doc.add_paragraph(b, style='List Bullet')

    add_heading_styled(doc, 'Step 2: Michigan Tax Tribunal (If Board Denies)', level=2)
    doc.add_paragraph(
        'If the Board does not grant adequate relief, file with the Michigan Tax Tribunal '
        'Small Claims Division by July 31, 2026. No filing fee if your property has a '
        'Principal Residence Exemption of at least 50%. File at michigan.gov/taxtrib.'
    )

    # Section 6: What to include
    add_heading_styled(doc, 'What to Include in Your Appeal', level=1)
    include_bullets = [
        'Completed Form L-4035 (download from michigan.gov or request from township)',
        'Your property record card - VERIFY that square footage, bedrooms, bathrooms, '
        'lot size, and condition rating are correct. Request from BSA Online (bsaonline.com/?uid=193)',
        'Comparable sales analysis (use the table above or customize with your own comps)',
        'The township\'s own ECF data showing AR-4 ECF of 0.802 (2026)',
        'Any property-specific issues: needed repairs, functional obsolescence, etc.',
        'Professional appraisal ($300-$500) for the strongest possible case',
    ]
    for b in include_bullets:
        doc.add_paragraph(b, style='List Bullet')

    # Section 7: Legal standards
    add_heading_styled(doc, 'Legal Standards That Favor Homeowners', level=1)
    doc.add_paragraph(
        'Under Michigan law (MCL 211.27), "true cash value" means the usual selling price - '
        'not replacement cost, not automated valuations (like Zillow Zestimates), and not '
        'listing prices. Key legal precedents:'
    )
    legal_bullets = [
        'There is NO presumption that the assessor\'s value is correct '
        '(Alhi Development Co v Orion Twp, 110 Mich App 764, 1981)',
        'The sales-comparison approach is the most persuasive method for residential property '
        '(Meadowlanes Ltd v Holland, 437 Mich 473, 1991)',
        '"True cash value" is the usual selling price in an arm\'s-length transaction '
        '(Huron Ridge LP v Ypsilanti Twp, 275 Mich App 23, 2007)',
    ]
    for b in legal_bullets:
        doc.add_paragraph(b, style='List Bullet')

    # Contact info
    doc.add_paragraph('_' * 75)
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = contact.add_run('Pittsfield Township Assessing Office\n')
    run.bold = True
    run.font.size = Pt(11)
    contact.add_run('6201 W. Michigan Avenue, Ann Arbor, MI 48108\n')
    contact.add_run('Phone: 734-822-3115 | Email: assessing@pittsfield-mi.gov\n')
    contact.add_run('BSA Online: bsaonline.com/?uid=193 | Township: pittsfield-mi.gov')

    doc.save(OUTPUT)
    print(f"Generated: {OUTPUT}")

if __name__ == "__main__":
    main()
