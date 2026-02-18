#!/usr/bin/env python3
"""Generate Property Tax Appeal Guide for 3543 Fieldcrest Ln as Word document."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

OUTPUT = os.path.join(os.path.dirname(__file__), "Property_Tax_Appeal_Guide_Fieldcrest_Ln.docx")


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)
    return h


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = val
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.size = Pt(9)
    return table


def main():
    doc = Document()

    # Default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # --- Title ---
    title = doc.add_heading('2026 Property Tax Appeal Guide', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('3543 Fieldcrest Ln | University Palisades (UNF) | Pittsfield Township')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(100, 100, 100)

    # Deadline
    deadline = doc.add_paragraph()
    deadline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = deadline.add_run(
        '\nDEADLINE: Written appeals must be received by March 10, 2026 at 5:00 PM\n'
    )
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(180, 0, 0)
    run2 = deadline.add_run(
        'Pittsfield Township Assessing, 6201 W. Michigan Ave, Ann Arbor, MI 48108\n'
    )
    run2.font.size = Pt(10)

    doc.add_paragraph('_' * 75)

    # --- Section 1: Your Property ---
    add_heading_styled(doc, 'Your Property Assessment', level=1)

    add_table(doc, ["Item", "Value"], [
        ["Property Address", "3543 Fieldcrest Ln, Ann Arbor, MI 48108"],
        ["Subdivision", "University Palisades (UNF)"],
        ["2026 SEV (State Equalized Value)", "$253,427"],
        ["2026 Assessed Value", "$253,427"],
        ["2026 Taxable Value", "$201,370"],
        ["Implied True Cash Value (SEV x 2)", "$506,854"],
    ])
    doc.add_paragraph()

    # --- Section 2: Why Appeal ---
    add_heading_styled(doc, 'Why You Should Appeal', level=1)

    doc.add_paragraph(
        'The 2026 assessment for 3543 Fieldcrest Ln implies a True Cash Value of $506,854. '
        'This is demonstrably above market value based on the township\'s own sales data:'
    )

    bullets = [
        'Your implied TCV of $506,854 exceeds EVERY arm\'s-length sale ever recorded in '
        'University Palisades. The highest sale was $470,000 (3745 Fieldcrest Ln, Jun 2022).',
        'The five Fieldcrest Ln arm\'s-length sales average $445,800 (median $440,000) -- '
        'your assessment is $61,054 (13.7%) above the average.',
        'The most recent Fieldcrest Ln arm\'s-length sale was $430,000 (3781 Fieldcrest Ln, '
        'Jul 2024) -- your assessment is $76,854 (17.9%) above this sale.',
        'The township\'s own ECF for 3578 Fieldcrest Ln is 0.856, meaning the cost approach '
        'overvalues that property by 14.4%.',
        'Across all 21 arm\'s-length sales in University Palisades, the average is $416,452 '
        'and the median is $425,000 -- both far below your implied TCV.',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    # --- Section 3: Recommended Value ---
    add_heading_styled(doc, 'Recommended Value', level=1)

    doc.add_paragraph(
        'Based on comparable sales evidence, the following reductions are supported:'
    )

    add_table(doc, ["Scenario", "Recommended TCV", "Recommended SEV", "Reduction"], [
        ["Conservative (Fieldcrest avg)", "$445,800", "$222,900", "$30,527"],
        ["Moderate (Fieldcrest median)", "$440,000", "$220,000", "$33,427"],
        ["Market-based (most recent sale)", "$430,000", "$215,000", "$38,427"],
    ])
    doc.add_paragraph()

    # --- Section 4: Comparable Sales ---
    add_heading_styled(doc, 'Comparable Sales Evidence', level=1)

    add_heading_styled(doc, 'Fieldcrest Ln Sales (Most Relevant)', level=2)
    doc.add_paragraph(
        'All sales on the same street as the subject property. Every sale is below '
        'the implied TCV of $506,854.'
    )

    add_table(doc,
        ["Address", "Sale Price", "Date", "Terms", "vs $506,854"],
        [
            ["3781 Fieldcrest Ln", "$430,000", "Jul 2024", "Arm's Length", "-$76,854 (15.2%)"],
            ["3793 Fieldcrest Ln", "$455,000", "Jun 2023", "Arm's Length", "-$51,854 (10.2%)"],
            ["3578 Fieldcrest Ln", "$440,000", "Aug 2023", "Arm's Length", "-$66,854 (13.2%)"],
            ["3838 Fieldcrest Ln", "$434,000", "Apr 2022", "Arm's Length", "-$72,854 (14.4%)"],
            ["3745 Fieldcrest Ln", "$470,000", "Jun 2022", "Arm's Length", "-$36,854 (7.3%)"],
            ["3533 Fieldcrest Ln*", "$350,000", "Feb 2025", "Non-arm's", "-$156,854 (30.9%)"],
        ],
    )
    doc.add_paragraph()
    p_note = doc.add_paragraph()
    run = p_note.add_run(
        '* 3533 Fieldcrest Ln was classified as "21-NOT USED/OTHER" by the township. '
        'If this was a legitimate market transaction, it further supports a lower value.'
    )
    run.font.size = Pt(9)
    run.italic = True

    add_heading_styled(doc, 'All University Palisades Arm\'s-Length Sales', level=2)
    doc.add_paragraph(
        'All 21 arm\'s-length sales in the subdivision. None reached $506,854.'
    )

    add_table(doc,
        ["Address", "Sale Price", "Date"],
        [
            ["3745 Fieldcrest Ln", "$470,000", "Jun 2022"],
            ["3899 Century Ct", "$462,500", "Feb 2024"],
            ["3793 Fieldcrest Ln", "$455,000", "Jun 2023"],
            ["3847 Palisades Blvd", "$454,000", "Apr 2024"],
            ["3750 Palisades Blvd", "$446,000", "Jul 2022"],
            ["3984 Palisades Blvd", "$440,000", "May 2023"],
            ["3578 Fieldcrest Ln", "$440,000", "Aug 2023"],
            ["3838 Fieldcrest Ln", "$434,000", "Apr 2022"],
            ["3781 Fieldcrest Ln", "$430,000", "Jul 2024"],
            ["3972 Palisades Blvd", "$430,000", "May 2024"],
            ["4538 Palisades Ct", "$425,000", "May 2022"],
            ["4542 Palisades Ct", "$420,000", "Jul 2022"],
            ["3720 Palisades Blvd", "$420,000", "Jan 2023"],
            ["3818 Century Ct", "$420,000", "Jul 2022"],
            ["3960 Palisades Blvd", "$400,000", "Aug 2023"],
            ["3979 Lancaster Ct", "$388,000", "Sep 2021"],
            ["3876 Palisades Blvd", "$385,000", "Apr 2021"],
            ["3859 Century Ct", "$380,000", "Oct 2022"],
            ["3855 Palisades Blvd", "$376,000", "Dec 2022"],
            ["3789 Palisades Blvd", "$375,000", "Jun 2021"],
            ["3936 Palisades Blvd", "$315,000", "Sep 2021"],
        ],
    )
    doc.add_paragraph()

    p_stats = doc.add_paragraph()
    p_stats.add_run('Subdivision Statistics: ').bold = True
    p_stats.add_run(
        'Average $416,452 | Median $425,000 | Range $315,000-$470,000 | '
        '0 of 21 sales reached $506,854'
    )

    # --- Section 5: ECF Data ---
    add_heading_styled(doc, 'Township ECF Data (Cost Approach vs. Market)', level=1)

    doc.add_paragraph(
        'The Economic Condition Factor (ECF) is the township\'s own metric comparing '
        'actual sale prices to cost-approach valuations. ECF below 1.0 = cost approach '
        'overvalues the property.'
    )

    add_table(doc,
        ["Address", "Sale Price", "ECF (2026)", "ECF (2025)"],
        [
            ["3578 Fieldcrest Ln", "$440,000", "0.856", "0.852"],
            ["3781 Fieldcrest Ln", "$430,000", "1.082", "--"],
            ["3793 Fieldcrest Ln", "$455,000", "1.033", "1.024"],
            ["3838 Fieldcrest Ln", "$434,000", "--", "1.004"],
            ["3745 Fieldcrest Ln", "$470,000", "--", "1.252"],
            ["3984 Palisades Blvd", "$440,000", "0.875", "0.869"],
            ["3960 Palisades Blvd", "$400,000", "1.123", "1.115"],
            ["3847 Palisades Blvd", "$454,000", "1.112", "--"],
            ["3899 Century Ct", "$462,500", "1.032", "1.106"],
            ["3855 Palisades Blvd", "$376,000", "--", "0.835"],
        ],
    )
    doc.add_paragraph()

    ecf_note = doc.add_paragraph()
    ecf_note.add_run('Key: ').bold = True
    ecf_note.add_run(
        '3578 Fieldcrest Ln (closest comparable) has an ECF of 0.856, meaning the '
        'cost approach overvalues it by 14.4%. 8 of 13 properties in the 2025 study '
        'had ECFs below 1.0.'
    )

    # --- Section 6: How Assessment Works ---
    add_heading_styled(doc, 'How Michigan Property Tax Assessment Works', level=1)
    doc.add_paragraph(
        'Under Michigan law, the Assessed Value (SEV) must equal 50% of the property\'s '
        'True Cash Value (TCV), defined as the "usual selling price" -- the price a willing '
        'buyer and seller would agree to in an arm\'s-length transaction (MCL 211.27). '
        'Your Taxable Value (TV) is capped by Proposal A at the lesser of 5% or the CPI '
        'increase per year.'
    )
    doc.add_paragraph(
        'Your current taxable value ($201,370) is below the SEV ($253,427). Even if the '
        'assessed value is reduced to $222,900, it would still be above the taxable value. '
        'However, the reduction constrains future TV growth and establishes a lower baseline '
        'if the property is transferred (TV resets to SEV upon sale). If the SEV is reduced '
        'below $201,370, your taxes would decrease immediately.'
    )

    # --- Section 7: Filing Instructions ---
    add_heading_styled(doc, 'How to File Your Appeal', level=1)

    add_heading_styled(doc, 'Step 1: Board of Review (Required First Step)', level=2)
    doc.add_paragraph(
        'You MUST file with the Pittsfield Township Board of Review first. '
        'Skipping this forfeits your right to appeal to the Michigan Tax Tribunal.'
    )
    step1_bullets = [
        'Public Hearings: March 9 (9am-12pm and 6pm-9pm), March 10 (9am-12pm), '
        'March 11 (1pm-5pm)',
        'Written Appeal Deadline: March 10, 2026 at 5:00 PM',
        'Location: 6201 W. Michigan Avenue, Ann Arbor, MI 48108',
        'Phone: 734-822-3115 | Email: assessing@pittsfield-mi.gov',
        'Form Required: Petition to Board of Review (Form L-4035)',
        'Written petitions are accepted in lieu of personal appearance',
        'If mailing, send certified with return receipt requested',
    ]
    for b in step1_bullets:
        doc.add_paragraph(b, style='List Bullet')

    add_heading_styled(doc, 'Step 2: Michigan Tax Tribunal (If Board Denies)', level=2)
    doc.add_paragraph(
        'If the Board does not grant adequate relief, file with the Michigan Tax Tribunal '
        'Small Claims Division by July 31, 2026. No filing fee if your property has a '
        'Principal Residence Exemption of at least 50%. File at michigan.gov/taxtrib.'
    )

    # --- Section 8: What to Include ---
    add_heading_styled(doc, 'What to Include in Your Appeal', level=1)
    include_bullets = [
        'Completed Form L-4035 (download from michigan.gov or request from township)',
        'Your property record card -- VERIFY that square footage, bedrooms, bathrooms, '
        'lot size, and condition rating are correct. Request from BSA Online '
        '(bsaonline.com/?uid=193). Any errors strengthen your case.',
        'This appeal guide with comparable sales analysis',
        'The township\'s own ECF data showing Fieldcrest Ln ECFs as low as 0.856',
        'Any property-specific issues: needed repairs, functional obsolescence, etc.',
        'Professional appraisal ($300-$500) for the strongest possible case',
    ]
    for b in include_bullets:
        doc.add_paragraph(b, style='List Bullet')

    # --- Section 9: Legal Standards ---
    add_heading_styled(doc, 'Legal Standards That Favor Homeowners', level=1)
    doc.add_paragraph(
        'Under Michigan law (MCL 211.27), "true cash value" means the usual selling price -- '
        'not replacement cost, not automated valuations, and not listing prices. '
        'Key legal precedents:'
    )
    legal_bullets = [
        'There is NO presumption that the assessor\'s value is correct '
        '(Alhi Development Co v Orion Twp, 110 Mich App 764, 1981)',
        'The sales-comparison approach is the most persuasive method for residential property '
        '(Meadowlanes Ltd v Holland, 437 Mich 473, 1991)',
        '"True cash value" is the usual selling price in an arm\'s-length transaction '
        '(Huron Ridge LP v Ypsilanti Twp, 275 Mich App 23, 2007)',
        'The Tax Tribunal has an independent duty to determine the correct value '
        '(Great Lakes Div of Nat\'l Steel v Ecorse, 227 Mich App 379, 1998)',
    ]
    for b in legal_bullets:
        doc.add_paragraph(b, style='List Bullet')

    # --- Contact Info ---
    doc.add_paragraph('_' * 75)
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = contact.add_run('Pittsfield Township Assessing Office\n')
    run.bold = True
    run.font.size = Pt(11)
    contact.add_run('6201 W. Michigan Avenue, Ann Arbor, MI 48108\n')
    contact.add_run('Phone: 734-822-3115 | Email: assessing@pittsfield-mi.gov\n')
    contact.add_run('BSA Online: bsaonline.com/?uid=193 | Township: pittsfield-mi.gov')

    doc.save(OUTPUT)
    print(f"Generated: {OUTPUT}")


if __name__ == "__main__":
    main()

